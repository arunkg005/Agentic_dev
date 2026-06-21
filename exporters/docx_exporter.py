import logging
# pyrefly: ignore [missing-import]
import docx
# pyrefly: ignore [missing-import]
from docx.shared import Pt, RGBColor

logger = logging.getLogger("NGO-Campaign-Copilot.DOCXExporter")

def compile_markdown_to_docx(markdown_content: str, output_path: str) -> bool:
    """Generates a professionally-styled Microsoft Word document (.docx) from Markdown."""
    try:
        doc = docx.Document()
        
        # Base Document Font Setup
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10.5)
        
        lines = markdown_content.split('\n')
        in_table = False
        table_rows = []
        
        def flush_table():
            nonlocal in_table, table_rows
            if not table_rows:
                in_table = False
                return
            # Determine number of columns
            cols = len(table_rows[0])
            table = doc.add_table(rows=len(table_rows), cols=cols)
            table.style = 'Light Shading Accent 1'
            for r_idx, row in enumerate(table_rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < len(table.rows[r_idx].cells):
                        table.rows[r_idx].cells[c_idx].text = cell_text.strip()
            table_rows = []
            in_table = False
            doc.add_paragraph() # Spacing

        for line in lines:
            stripped = line.strip()
            
            # Table handling
            if stripped.startswith('|'):
                if '---' in stripped or ':---' in stripped:
                    continue # Skip separator row
                in_table = True
                # Extract cell values
                parts = [p.strip() for p in stripped.split('|')[1:-1]]
                table_rows.append(parts)
                continue
            else:
                if in_table:
                    flush_table()
            
            if not stripped:
                continue
                
            # Headers
            if stripped.startswith('# '):
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(18)
                h.paragraph_format.space_after = Pt(6)
                h.paragraph_format.keep_with_next = True
                run = h.add_run(stripped[2:])
                run.font.name = 'Arial'
                run.font.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(15, 23, 42) # Slate #0f172a
            elif stripped.startswith('## '):
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(14)
                h.paragraph_format.space_after = Pt(4)
                h.paragraph_format.keep_with_next = True
                run = h.add_run(stripped[3:])
                run.font.name = 'Arial'
                run.font.bold = True
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(13, 148, 136) # Teal #0d9488
            elif stripped.startswith('### '):
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(4)
                h.paragraph_format.keep_with_next = True
                run = h.add_run(stripped[4:])
                run.font.name = 'Arial'
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(30, 41, 59) # Dark Slate #1e293b
            elif stripped.startswith('- ') or stripped.startswith('* '):
                text = stripped[2:]
                p_item = doc.add_paragraph(style='List Bullet')
                # Parse bold markdown **
                if '**' in text:
                    parts = text.split('**')
                    for i, part in enumerate(parts):
                        run = p_item.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
                else:
                    p_item.add_run(text)
            elif stripped.startswith('---'):
                # Line break
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            else:
                p_para = doc.add_paragraph()
                # Parse bold markdown **
                if '**' in stripped:
                    parts = stripped.split('**')
                    for i, part in enumerate(parts):
                        run = p_para.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
                else:
                    p_para.add_run(stripped)
                    
        if in_table:
            flush_table()
            
        doc.save(output_path)
        logger.info("Campaign DOCX saved to '%s'.", output_path)
        return True
    except Exception as e:
        logger.error("Error generating DOCX: %s", e)
        return False
