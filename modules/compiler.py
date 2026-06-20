import os
import logging
import markdown
import docx
from docx.shared import Pt, RGBColor
from xhtml2pdf import pisa
from config import call_llm

logger = logging.getLogger("NGO-Campaign-Copilot.Compiler")


def generate_strategic_recommendations(
    campaign_summary: str,
    api_key: str | None = None, provider: str = "Gemini (Google)",
) -> str:
    """
    Calls the LLM to produce 3–4 high-level strategic recommendations
    tailored to the campaign context.
    """
    logger.info("Generating strategic recommendations...")
    prompt = (
        "You are a senior advisor for a large international NGO. "
        "Review this campaign plan summary:\n\n"
        f"{campaign_summary}\n\n"
        "Provide 3 to 4 high-value, practical strategic recommendations "
        "to ensure the campaign's success. Do not repeat details already in the summary. "
        "Write them as a simple markdown bullet list."
    )
    try:
        return call_llm(prompt, api_key=api_key, json_output=False, provider=provider)
    except Exception as e:
        logger.error("Error generating recommendations: %s", e)
        return (
            "- **Focus on Local Engagement**: Build relationships with local schools or community centres early.\n"
            "- **Optimise Communication**: Maintain a shared WhatsApp group for real-time volunteer coordination.\n"
            "- **Document for Future**: Take photos/videos to demonstrate impact for future donor proposals."
        )


def compile_campaign_plan(
    name: str,
    topic: str,
    audience: str,
    duration: int,
    duration_unit: str,
    budget: float,
    reach: int,
    objectives: list[str],
    timeline: list[dict],
    resources: dict,
    volunteers: dict,
    risks: list[dict],
    metrics: list[dict],
    api_key: str | None = None,
    provider: str = "Gemini (Google)",
) -> tuple[str, str]:
    """
    Compiles every generated component into a single, professional Markdown report.
    Saves the file to the ``output/`` directory and returns (markdown_content, file_path).
    """
    logger.info("Compiling campaign plan for '%s'...", name)

    # --- Build markdown segments with safe .get() everywhere ---
    objectives_md = "\n".join(f"* {obj}" for obj in objectives)

    timeline_md = ""
    for idx, t in enumerate(timeline, 1):
        timeline_md += f"### {t.get('phase', f'Phase {idx}')} ({t.get('timeframe', '')})\n"
        for act in t.get("activities", []):
            timeline_md += f"- {act}\n"
        timeline_md += "\n"

    phys_md = "\n".join(
        f"| {p.get('item', '')} | {p.get('suggested_quantity', '')} | {p.get('purpose', '')} |"
        for p in resources.get("physical_materials", [])
    )
    digi_md = "\n".join(
        f"| {d.get('item', '')} | {d.get('purpose', '')} |"
        for d in resources.get("digital_resources", [])
    )
    budget_md = "\n".join(
        f"| {b.get('category', '')} | {b.get('percentage', 0)}% | INR {b.get('amount_inr', 0)} |"
        for b in resources.get("budget_allocation", [])
    )
    vol_md = "\n".join(
        f"| {r.get('role_name', '')} | {r.get('count', 0)} | {r.get('responsibilities', '')} |"
        for r in volunteers.get("roles", [])
    )
    risks_md = "\n".join(
        f"| {rk.get('risk', '')} | {rk.get('mitigation', '')} |"
        for rk in risks
    )
    metrics_md = "\n".join(
        f"| {m.get('name', '')} | {m.get('target', '')} | {m.get('method', '')} |"
        for m in metrics
    )

    vol_count = resources.get("calculated_volunteers", 0)

    # Strategic recommendations
    summary_ctx = (
        f"Campaign: {name}\nTopic: {topic}\nAudience: {audience}\n"
        f"Budget: INR {budget}\nReach: {reach}\n"
        f"Objectives: {', '.join(objectives[:2])}\nVolunteers: {vol_count}"
    )
    recommendations_md = generate_strategic_recommendations(summary_ctx, api_key=api_key, provider=provider)

    # --- Assemble final document ---
    markdown_content = f"""# NGO Campaign Execution Plan: {name}

## Campaign Overview
- **Topic / Focus**: {topic}
- **Target Audience**: {audience}
- **Campaign Duration**: {duration} {duration_unit}
- **Expected Reach**: {reach} participants
- **Projected Budget**: INR {budget}
- **Total Volunteer Workforce**: {vol_count} volunteers

---

## Objectives
{objectives_md}

---

## Execution Timeline

{timeline_md}
---

## Resource Requirements

### Physical Materials
| Material Item | Quantity | Purpose |
| :--- | :--- | :--- |
{phys_md}

### Digital Assets & Platforms
| Tool / Platform | Purpose |
| :--- | :--- |
{digi_md}

### Budget Allocation
| Expense Category | Allocation % | Estimated Cost |
| :--- | :---: | :---: |
{budget_md}

*Note: All cost estimates are preliminary and subject to local vendor quotes.*

---

## Volunteer Mobilization Plan
Distribution of roles for {vol_count} volunteers:

| Volunteer Role | Count | Core Responsibilities |
| :--- | :---: | :--- |
{vol_md}

**Coordination Tip**: *{volunteers.get('coordination_tips', 'N/A')}*

---

## Risk & Mitigation Assessment

| Identified Risk | Mitigation Strategy |
| :--- | :--- |
{risks_md}

---

## Success Metrics & Evaluation

| Metric | Target Value | Measurement Method |
| :--- | :--- | :--- |
{metrics_md}

---

## Strategic Recommendations
{recommendations_md}

---
*Document generated by Naye Pankh NGO Campaign Planning Copilot.*
"""

    # Save to file
    os.makedirs("output", exist_ok=True)
    safe_name = "".join(c if c.isalnum() else "_" for c in name.strip()).lower()
    # Guard against empty filenames
    if not safe_name:
        safe_name = "campaign"
    md_file_path = os.path.join("output", f"campaign_plan_{safe_name}.md")
    pdf_file_path = os.path.join("output", f"campaign_plan_{safe_name}.pdf")
    docx_file_path = os.path.join("output", f"campaign_plan_{safe_name}.docx")

    # 1. Save Markdown
    with open(md_file_path, "w", encoding="utf-8") as f:
        f.write(markdown_content)
    logger.info("Campaign plan saved to '%s'.", md_file_path)

    # 2. Save PDF
    html_body = markdown.markdown(markdown_content, extensions=['tables'])
    compile_campaign_plan_pdf(html_body, pdf_file_path)

    # 3. Save DOCX
    compile_markdown_to_docx(markdown_content, docx_file_path)

    return markdown_content, md_file_path, pdf_file_path, docx_file_path


def compile_campaign_plan_pdf(html_content: str, output_path: str) -> bool:
    """Converts the campaign HTML preview to a beautifully formatted PDF."""
    styled_html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
    @page {{
        size: letter;
        margin: 0.8in;
        @bottom-right {{
            content: "Page " counter(page);
            font-family: Arial, sans-serif;
            font-size: 8pt;
            color: #64748b;
        }}
    }}
    body {{
        font-family: Arial, sans-serif;
        color: #1e293b;
        line-height: 1.5;
        font-size: 10pt;
    }}
    h1 {{
        color: #0f172a;
        font-size: 18pt;
        border-bottom: 2px solid #0d9488;
        padding-bottom: 6px;
        margin-top: 0;
        margin-bottom: 12px;
    }}
    h2 {{
        color: #0d9488;
        font-size: 13pt;
        margin-top: 18px;
        margin-bottom: 8px;
        border-bottom: 1px solid #cbd5e1;
        padding-bottom: 4px;
        page-break-after: avoid;
    }}
    h3 {{
        color: #1e293b;
        font-size: 10.5pt;
        margin-top: 14px;
        margin-bottom: 6px;
        page-break-after: avoid;
    }}
    p, li {{
        font-size: 9.5pt;
        margin-bottom: 6px;
    }}
    table {{
        width: 100%;
        border-collapse: collapse;
        margin-top: 8px;
        margin-bottom: 12px;
        font-size: 8.5pt;
    }}
    th, td {{
        border: 1px solid #e2e8f0;
        padding: 6px 8px;
        text-align: left;
    }}
    th {{
        background-color: #f1f5f9;
        font-weight: bold;
        color: #0f172a;
        border-bottom: 2px solid #cbd5e1;
    }}
    ul, ol {{
        margin-top: 4px;
        margin-bottom: 8px;
        padding-left: 20px;
    }}
    hr {{
        border: none;
        border-top: 1px solid #e2e8f0;
        margin: 16px 0;
    }}
</style>
</head>
<body>
{html_content}
</body>
</html>
"""
    try:
        with open(output_path, "w+b") as f:
            pisa_status = pisa.CreatePDF(styled_html, dest_file=f)
        if pisa_status.err:
            logger.error("xhtml2pdf error code: %s", pisa_status.err)
            return False
        logger.info("Campaign PDF saved to '%s'.", output_path)
        return True
    except Exception as e:
        logger.error("Error generating PDF: %s", e)
        return False


def compile_markdown_to_docx(markdown_content: str, output_path: str) -> bool:
    """Generates a professionally-styled Microsoft Word document (.docx) from Markdown."""
    try:
        doc = docx.Document()
        
        # Base Document Font Setup
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10.5)
        
        lines = markdown_content.split('\n')
        in_table = False
        table_rows = []
        
        def flush_table():
            nonlocal in_table, table_rows
            if not table_rows:
                in_table = False
                return
            # Determine number of columns
            cols = len(table_rows[0])
            table = doc.add_table(rows=len(table_rows), cols=cols)
            table.style = 'Light Shading Accent 1'
            for r_idx, row in enumerate(table_rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx < len(table.rows[r_idx].cells):
                        table.rows[r_idx].cells[c_idx].text = cell_text.strip()
            table_rows = []
            in_table = False
            doc.add_paragraph() # Spacing

        for line in lines:
            stripped = line.strip()
            
            # Table handling
            if stripped.startswith('|'):
                if '---' in stripped or ':---' in stripped:
                    continue # Skip separator row
                in_table = True
                # Extract cell values
                parts = [p.strip() for p in stripped.split('|')[1:-1]]
                table_rows.append(parts)
                continue
            else:
                if in_table:
                    flush_table()
            
            if not stripped:
                continue
                
            # Headers
            if stripped.startswith('# '):
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(18)
                h.paragraph_format.space_after = Pt(6)
                h.paragraph_format.keep_with_next = True
                run = h.add_run(stripped[2:])
                run.font.name = 'Arial'
                run.font.bold = True
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(15, 23, 42) # Slate #0f172a
            elif stripped.startswith('## '):
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(14)
                h.paragraph_format.space_after = Pt(4)
                h.paragraph_format.keep_with_next = True
                run = h.add_run(stripped[3:])
                run.font.name = 'Arial'
                run.font.bold = True
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(13, 148, 136) # Teal #0d9488
            elif stripped.startswith('### '):
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(4)
                h.paragraph_format.keep_with_next = True
                run = h.add_run(stripped[4:])
                run.font.name = 'Arial'
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(30, 41, 59) # Dark Slate #1e293b
            elif stripped.startswith('- ') or stripped.startswith('* '):
                text = stripped[2:]
                p_item = doc.add_paragraph(style='List Bullet')
                # Parse bold markdown **
                if '**' in text:
                    parts = text.split('**')
                    for i, part in enumerate(parts):
                        run = p_item.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
                else:
                    p_item.add_run(text)
            elif stripped.startswith('---'):
                # Line break
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            else:
                p_para = doc.add_paragraph()
                # Parse bold markdown **
                if '**' in stripped:
                    parts = stripped.split('**')
                    for i, part in enumerate(parts):
                        run = p_para.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
                else:
                    p_para.add_run(stripped)
                    
        if in_table:
            flush_table()
            
        doc.save(output_path)
        logger.info("Campaign DOCX saved to '%s'.", output_path)
        return True
    except Exception as e:
        logger.error("Error generating DOCX: %s", e)
        return False


def refine_campaign_plan(
    current_markdown: str,
    prompt: str,
    api_key: str | None = None,
    provider: str = "Gemini (Google)"
) -> tuple[str, str, str, str]:
    """
    Calls the LLM to refine an existing campaign plan based on a prompt.
    Regenerates Markdown, HTML, PDF, and DOCX files.
    Returns (refined_markdown, md_file_path, pdf_file_path, docx_file_path).
    """
    logger.info("Refining campaign plan with prompt: %s", prompt)
    
    # Construct prompt
    llm_prompt = (
        "You are an expert AI Campaign Planning Assistant.\n"
        "Here is the current campaign execution plan (in Markdown format):\n\n"
        f"{current_markdown}\n\n"
        "The user wants to refine/update this plan with the following request:\n"
        f"\"{prompt}\"\n\n"
        "Please provide the updated, complete campaign plan in Markdown format. "
        "Keep the exact same layout, header structure, table formats, and styling. "
        "Modify ONLY the sections that need changes based on the user's request. "
        "Output ONLY the markdown content. Do not include any introductory or concluding comments."
    )
    
    try:
        refined_markdown = call_llm(llm_prompt, api_key=api_key, json_output=False, provider=provider)
    except Exception as e:
        logger.error("Error calling LLM for refinement: %s", e)
        refined_markdown = current_markdown + f"\n\n*Refinement failed: {str(e)}*"

    # Extract the campaign name to construct a safe filename
    campaign_name = "refined_campaign"
    for line in refined_markdown.split("\n")[:5]:
        if line.startswith("# NGO Campaign Execution Plan:"):
            campaign_name = line.replace("# NGO Campaign Execution Plan:", "").strip()
            break
            
    safe_name = "".join(c if c.isalnum() else "_" for c in campaign_name.strip()).lower()
    if not safe_name:
        safe_name = "campaign"
        
    os.makedirs("output", exist_ok=True)
    md_file_path = os.path.join("output", f"campaign_plan_{safe_name}.md")
    pdf_file_path = os.path.join("output", f"campaign_plan_{safe_name}.pdf")
    docx_file_path = os.path.join("output", f"campaign_plan_{safe_name}.docx")

    # 1. Save Markdown
    with open(md_file_path, "w", encoding="utf-8") as f:
        f.write(refined_markdown)

    # 2. Save PDF
    html_body = markdown.markdown(refined_markdown, extensions=['tables'])
    compile_campaign_plan_pdf(html_body, pdf_file_path)

    # 3. Save DOCX
    compile_markdown_to_docx(refined_markdown, docx_file_path)

    return refined_markdown, md_file_path, pdf_file_path, docx_file_path
